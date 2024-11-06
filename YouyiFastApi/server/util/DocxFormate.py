"""
 @Project: service
 @File: DocxFormate.py
 @Version: 1.0.0
 @Author: 轩祎SheyZhang
 @Created: 2024/5/29
 @Modified: 2024/5/29
 @Description: DocxFormate.py
"""
import os

from server.business.model.meeting import Meeting


class DocxFormate:
    def __init__(self, data: Meeting = None):
        self.doc_template = None
        self.data = data
        self.content = ''
        self.folder_path = 'C:\\fa'

    def format(self):
        self.content = {
            'meeting_id': self.data.meeting_id,
            'meeting_name': self.data.meeting_name,
            'meeting_tags': self.data.meeting_tags,
            'meeting_time': self.data.meeting_time,
            'meeting_location': self.data.meeting_location,
            'meeting_status': self.data.meeting_status,
            'meeting_host': self.data.meeting_host,
            'meeting_attendees': self.data.meeting_attendees,
            'meeting_description': self.data.meeting_description,
            'meeting_outline': self.data.meeting_outline,
            'meeting_subject': self.data.meeting_subject,
            'meeting_content': self.data.meeting_content,
            'meeting_duration': self.data.meeting_duration,
            'meeting_created_at': self.data.meeting_created_at,
            'meeting_updated_at': self.data.meeting_updated_at,
            'meeting_summary': {
                'summary_title': 'summary_title',
                'summary_content': 'summary_content',  # todo 纪要的内容提取
            }
        }

        return self.content

    def write(self, data):
        print(data)
        from docx import Document
        self.doc = Document('server/data/template/002.docx')

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.replace_text_in_paragraph(paragraph, data)

        for paragraph in self.doc.paragraphs:
            if paragraph not in self.doc.tables:
                self.replace_text_in_paragraph(paragraph, data)

    def replace_text_in_paragraph(self, paragraph, data):
        for key, value in data.items():
            placeholder = '{' + key + '}'
            if placeholder == '{meeting_summary}':
                paragraph.text = paragraph.text.replace(placeholder, str(data['meeting_summary']['summary_content']))
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    def save(self, data):
        path = os.path.join(self.folder_path, data['meeting_name'] + ".docx")
        self.doc.save(path)
